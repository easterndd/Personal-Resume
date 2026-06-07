import json
import os
from datetime import datetime
from typing import Dict, Any

import fitz  # PyMuPDF
from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..config import settings


class ExportService:
    def __init__(self):
        self.output_dir = settings.output_dir
        self.template_dir = os.path.join(os.path.dirname(__file__), "../templates")
        self.env = Environment(loader=FileSystemLoader(self.template_dir))
        os.makedirs(self.output_dir, exist_ok=True)

    def format_date(self, date_str: str) -> str:
        if not date_str:
            return ""
        parts = date_str.split("-")
        if len(parts) >= 2:
            return f"{parts[0]}.{parts[1].zfill(2)}"
        return date_str

    def format_date_range(self, start_date: str, end_date: str) -> str:
        start = self.format_date(start_date)
        end = self.format_date(end_date) if end_date else "至今"
        return f"{start} - {end}"

    def render_html(self, resume_data: Dict[str, Any], template_id: str = "modern") -> str:
        template = self.env.get_template(f"{template_id}/resume.html")
        return template.render(
            data=resume_data,
            format_date=self.format_date,
            format_date_range=self.format_date_range,
        )

    async def export_pdf(self, resume_data: Dict[str, Any], resume_id: str, template_id: str = "modern") -> str:
        """使用PyMuPDF生成PDF"""
        try:
            filename = f"{resume_data.get('basics', {}).get('name', 'resume')}_{resume_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)
            
            blue = (0.145, 0.388, 0.922)
            dark = (0.118, 0.161, 0.231)
            gray = (0.392, 0.455, 0.518)
            text_gray = (0.278, 0.333, 0.412)
            light_gray = (0.886, 0.906, 0.941)
            
            margin = 56  # 20mm
            y = margin
            
            basics = resume_data.get("basics", {})
            name = basics.get("name", "简历")
            headline = basics.get("headline", "")
            
            avatar = basics.get("avatar", "")
            has_avatar = avatar and avatar.startswith("data:image/")
            
            if has_avatar:
                try:
                    import base64
                    img_data = avatar.split(",")[1]
                    img_bytes = base64.b64decode(img_data)
                    pix = fitz.Pixmap(img_bytes)
                    img_width = pix.width
                    img_height = pix.height
                    max_dim = 80
                    scale = min(max_dim / img_width, max_dim / img_height)
                    img_width = int(img_width * scale)
                    img_height = int(img_height * scale)
                    avatar_x = 595 - margin - img_width
                    avatar_y = y
                    page.insert_image(fitz.Rect(avatar_x, avatar_y, avatar_x + img_width, avatar_y + img_height), stream=img_bytes)
                except:
                    has_avatar = False
            
            text_width = 595 - margin * 2 - (80 if has_avatar else 0)
            
            fontname = "china-s"
            try:
                test_font = fitz.Font(fontname)
            except:
                fontname = "helv"
            
            text_y = y
            rect = fitz.Rect(margin, text_y, margin + text_width, text_y + 32)
            page.insert_textbox(rect, name, fontsize=24, fontname=fontname, color=dark, align=0)
            y += 35
            
            if headline:
                rect = fitz.Rect(margin, y, margin + text_width, y + 18)
                page.insert_textbox(rect, headline, fontsize=12, fontname=fontname, color=blue, align=0)
                y += 20
            
            contact_parts = []
            if basics.get("phone"):
                contact_parts.append(f"电话：{basics['phone']}")
            if basics.get("email"):
                contact_parts.append(f"邮箱：{basics['email']}")
            if basics.get("location"):
                contact_parts.append(f"地点：{basics['location']}")
            if basics.get("github"):
                github = basics['github']
                if len(github) > 30:
                    github = github[:27] + "..."
                contact_parts.append(f"GitHub：{github}")
            
            if contact_parts:
                contact_text = " | ".join(contact_parts)
                rect = fitz.Rect(margin, y, margin + text_width, y + 16)
                page.insert_textbox(rect, contact_text, fontsize=9, fontname=fontname, color=gray, align=0)
                y += 22
            
            page.draw_line(fitz.Point(margin, y), fitz.Point(595 - margin, y), color=blue, width=2)
            y += 20
            
            def add_section_title(title, y_pos):
                page.draw_line(fitz.Point(margin, y_pos), fitz.Point(margin, y_pos + 14), color=blue, width=3)
                rect = fitz.Rect(margin + 10, y_pos, 595 - margin, y_pos + 16)
                page.insert_textbox(rect, title, fontsize=11, fontname=fontname, color=dark)
                page.draw_line(fitz.Point(margin + 10 + len(title) * 11 + 8, y_pos + 7), fitz.Point(595 - margin, y_pos + 7), color=light_gray, width=1)
                return y_pos + 20
            
            summary = resume_data.get("summary", "")
            if summary:
                y = add_section_title("个人简介", y)
                lines = self._wrap_text(summary, 60)
                for line in lines[:10]:
                    rect = fitz.Rect(margin, y, 595 - margin, y + 15)
                    page.insert_textbox(rect, line, fontsize=10, fontname=fontname, color=text_gray)
                    y += 14
                y += 8
            
            work_list = resume_data.get("work", [])
            if work_list:
                y = add_section_title("工作经历", y)
                for work in work_list:
                    company = work.get("company", "")
                    position = work.get("position", "")
                    location = work.get("location", "")
                    
                    company_line = company
                    if position:
                        company_line += f" · {position}"
                    if location:
                        company_line += f" | {location}"
                    
                    rect = fitz.Rect(margin, y, 595 - margin - 70, y + 16)
                    page.insert_textbox(rect, company_line, fontsize=10, fontname=fontname, color=dark)
                    
                    date_range = self.format_date_range(work.get("start_date", ""), work.get("end_date", ""))
                    rect = fitz.Rect(595 - margin - 70, y, 595 - margin, y + 14)
                    page.insert_textbox(rect, date_range, fontsize=9, fontname=fontname, color=blue)
                    y += 18
                    
                    if work.get("description"):
                        desc_lines = self._wrap_text(work["description"], 65)
                        for line in desc_lines[:5]:
                            rect = fitz.Rect(margin, y, 595 - margin, y + 14)
                            page.insert_textbox(rect, line, fontsize=9, fontname=fontname, color=text_gray)
                            y += 13
                        y += 4
                    
                    for highlight in work.get("highlights", []):
                        highlight_lines = self._wrap_text(f"• {highlight}", 62)
                        for line in highlight_lines[:2]:
                            rect = fitz.Rect(margin + 12, y, 595 - margin, y + 14)
                            page.insert_textbox(rect, line, fontsize=9, fontname=fontname, color=text_gray)
                            y += 13
                    y += 10
            
            projects = resume_data.get("projects", [])
            if projects:
                y = add_section_title("项目经历", y)
                for project in projects:
                    proj_name = project.get("name", "")
                    role = project.get("role", "")
                    
                    name_line = proj_name
                    if role:
                        name_line += f" · {role}"
                    
                    rect = fitz.Rect(margin, y, 595 - margin - 70, y + 16)
                    page.insert_textbox(rect, name_line, fontsize=10, fontname=fontname, color=dark)
                    
                    date_range = self.format_date_range(project.get("start_date", ""), project.get("end_date", ""))
                    rect = fitz.Rect(595 - margin - 70, y, 595 - margin, y + 14)
                    page.insert_textbox(rect, date_range, fontsize=9, fontname=fontname, color=blue)
                    y += 18
                    
                    if project.get("description"):
                        desc_lines = self._wrap_text(project["description"], 65)
                        for line in desc_lines[:5]:
                            rect = fitz.Rect(margin, y, 595 - margin, y + 14)
                            page.insert_textbox(rect, line, fontsize=9, fontname=fontname, color=text_gray)
                            y += 13
                        y += 4
                    
                    for highlight in project.get("highlights", []):
                        highlight_lines = self._wrap_text(f"• {highlight}", 62)
                        for line in highlight_lines[:2]:
                            rect = fitz.Rect(margin + 12, y, 595 - margin, y + 14)
                            page.insert_textbox(rect, line, fontsize=9, fontname=fontname, color=text_gray)
                            y += 13
                    
                    if project.get("technologies"):
                        tech_text = "技术栈：" + ", ".join(project["technologies"])
                        tech_lines = self._wrap_text(tech_text, 65)
                        for line in tech_lines[:2]:
                            rect = fitz.Rect(margin, y, 595 - margin, y + 14)
                            page.insert_textbox(rect, line, fontsize=9, fontname=fontname, color=gray)
                            y += 13
                    y += 10
            
            education = resume_data.get("education", [])
            if education:
                y = add_section_title("教育背景", y)
                for edu in education:
                    school = edu.get("school", "")
                    degree = edu.get("degree", "")
                    major = edu.get("major", "")
                    
                    edu_line = school
                    if degree:
                        edu_line += f" · {degree}"
                    if major:
                        edu_line += f" · {major}"
                    
                    rect = fitz.Rect(margin, y, 595 - margin - 70, y + 16)
                    page.insert_textbox(rect, edu_line, fontsize=10, fontname=fontname, color=dark)
                    
                    date_range = self.format_date_range(edu.get("start_date", ""), edu.get("end_date", ""))
                    rect = fitz.Rect(595 - margin - 70, y, 595 - margin, y + 14)
                    page.insert_textbox(rect, date_range, fontsize=9, fontname=fontname, color=blue)
                    y += 18
            
            skills = resume_data.get("skills", [])
            if skills:
                y = add_section_title("专业技能", y)
                for skill in skills:
                    category = skill.get("category", "")
                    items = ", ".join(skill.get("items", []))
                    skill_text = f"{category}：{items}"
                    skill_lines = self._wrap_text(skill_text, 65)
                    for line in skill_lines[:2]:
                        rect = fitz.Rect(margin, y, 595 - margin, y + 14)
                        page.insert_textbox(rect, line, fontsize=10, fontname=fontname, color=text_gray)
                        y += 13
                y += 8
            
            certificates = resume_data.get("certificates", [])
            if certificates:
                y = add_section_title("证书", y)
                for cert in certificates:
                    cert_name = cert.get("name", "")
                    cert_date = cert.get("date", "")
                    rect = fitz.Rect(margin, y, 595 - margin - 60, y + 14)
                    page.insert_textbox(rect, cert_name, fontsize=10, fontname=fontname, color=dark)
                    if cert_date:
                        rect = fitz.Rect(595 - margin - 60, y, 595 - margin, y + 14)
                        page.insert_textbox(rect, cert_date, fontsize=9, fontname=fontname, color=gray)
                    y += 16
            
            languages = resume_data.get("languages", [])
            if languages:
                y = add_section_title("语言能力", y)
                for lang in languages:
                    lang_name = lang.get("name", "")
                    lang_level = lang.get("level", "")
                    rect = fitz.Rect(margin, y, 595 - margin, y + 14)
                    page.insert_textbox(rect, f"{lang_name} · {lang_level}", fontsize=10, fontname=fontname, color=text_gray)
                    y += 14
            
            awards = resume_data.get("awards", [])
            if awards:
                y = add_section_title("荣誉奖项", y)
                for award in awards:
                    award_name = award.get("name", "")
                    award_date = award.get("date", "")
                    rect = fitz.Rect(margin, y, 595 - margin - 60, y + 14)
                    page.insert_textbox(rect, award_name, fontsize=10, fontname=fontname, color=dark)
                    if award_date:
                        rect = fitz.Rect(595 - margin - 60, y, 595 - margin, y + 14)
                        page.insert_textbox(rect, award_date, fontsize=9, fontname=fontname, color=gray)
                    y += 16
            
            doc.save(filepath)
            doc.close()
            
            return filepath
        except Exception as e:
            raise Exception(f"PDF export failed: {str(e)}")
    
    def _wrap_text(self, text: str, max_chars: int) -> list:
        """自动换行文本"""
        lines = []
        current_line = ""
        for char in text:
            if len(current_line) >= max_chars:
                lines.append(current_line)
                current_line = char
            else:
                current_line += char
        if current_line:
            lines.append(current_line)
        return lines

    def export_docx(self, resume_data: Dict[str, Any], resume_id: str) -> str:
        doc = Document()

        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        basics = resume_data.get("basics", {})
        title = doc.add_heading(basics.get("name", "简历"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.bold = True
        title.runs[0].font.color.rgb = RGBColor(37, 99, 235)

        if basics.get("headline"):
            headline = doc.add_paragraph(basics.get("headline"))
            headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
            headline.runs[0].font.size = Pt(12)

        contact_info = []
        if basics.get("phone"):
            contact_info.append(f"电话：{basics['phone']}")
        if basics.get("email"):
            contact_info.append(f"邮箱：{basics['email']}")
        if basics.get("location"):
            contact_info.append(f"地点：{basics['location']}")
        if basics.get("github"):
            contact_info.append(f"GitHub：{basics['github']}")
        if basics.get("website"):
            contact_info.append(f"网站：{basics['website']}")

        if contact_info:
            contact_para = doc.add_paragraph(" | ".join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()

        if resume_data.get("summary"):
            doc.add_heading("个人简介", level=1)
            summary_para = doc.add_paragraph(resume_data["summary"])
            summary_para.runs[0].font.size = Pt(10)

        for work in resume_data.get("work", []):
            doc.add_heading("工作经历", level=1)
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{work.get('company', '')}")
            company_run.bold = True
            company_run.font.size = Pt(10)
            if work.get("position"):
                company_para.add_run(f" - {work['position']}")
            if work.get("location"):
                company_para.add_run(f" | {work['location']}")

            date_para = doc.add_paragraph(self.format_date_range(work.get("start_date", ""), work.get("end_date", "")))
            date_para.runs[0].font.size = Pt(9)

            if work.get("description"):
                desc_para = doc.add_paragraph(work["description"])
                desc_para.runs[0].font.size = Pt(10)

            for highlight in work.get("highlights", []):
                highlight_para = doc.add_paragraph(f"• {highlight}", style="List Bullet")
                highlight_para.runs[0].font.size = Pt(10)

        for edu in resume_data.get("education", []):
            doc.add_heading("教育背景", level=1)
            school_para = doc.add_paragraph()
            school_run = school_para.add_run(f"{edu.get('school', '')}")
            school_run.bold = True
            school_run.font.size = Pt(10)
            if edu.get("degree"):
                school_para.add_run(f" - {edu['degree']}")
            if edu.get("major"):
                school_para.add_run(f" | {edu['major']}")

            date_para = doc.add_paragraph(self.format_date_range(edu.get("start_date", ""), edu.get("end_date", "")))
            date_para.runs[0].font.size = Pt(9)

            for highlight in edu.get("highlights", []):
                highlight_para = doc.add_paragraph(f"• {highlight}", style="List Bullet")
                highlight_para.runs[0].font.size = Pt(10)

        for project in resume_data.get("projects", []):
            doc.add_heading("项目经历", level=1)
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(f"{project.get('name', '')}")
            name_run.bold = True
            name_run.font.size = Pt(10)
            if project.get("role"):
                name_para.add_run(f" - {project['role']}")

            date_para = doc.add_paragraph(self.format_date_range(project.get("start_date", ""), project.get("end_date", "")))
            date_para.runs[0].font.size = Pt(9)

            if project.get("description"):
                desc_para = doc.add_paragraph(project["description"])
                desc_para.runs[0].font.size = Pt(10)

            for highlight in project.get("highlights", []):
                highlight_para = doc.add_paragraph(f"• {highlight}", style="List Bullet")
                highlight_para.runs[0].font.size = Pt(10)

            if project.get("technologies"):
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run("技术栈：")
                tech_run.bold = True
                tech_run.font.size = Pt(9)
                tech_para.add_run(", ".join(project["technologies"]))
                tech_para.runs[1].font.size = Pt(9)

        if resume_data.get("skills"):
            doc.add_heading("专业技能", level=1)
            for skill_category in resume_data["skills"]:
                skill_para = doc.add_paragraph()
                cat_run = skill_para.add_run(f"{skill_category.get('category', '')}：")
                cat_run.bold = True
                cat_run.font.size = Pt(10)
                skill_para.add_run(", ".join(skill_category.get("items", [])))
                skill_para.runs[1].font.size = Pt(10)

        filename = f"{basics.get('name', 'resume')}_{resume_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)

        return filepath

    def export_txt(self, resume_data: Dict[str, Any], resume_id: str) -> str:
        lines = []
        basics = resume_data.get("basics", {})

        lines.append(basics.get("name", "简历"))
        if basics.get("headline"):
            lines.append(basics.get("headline"))
        lines.append("")

        contact_info = []
        if basics.get("phone"):
            contact_info.append(f"电话：{basics['phone']}")
        if basics.get("email"):
            contact_info.append(f"邮箱：{basics['email']}")
        if basics.get("location"):
            contact_info.append(f"地点：{basics['location']}")
        lines.append(" | ".join(contact_info))
        lines.append("")

        if resume_data.get("summary"):
            lines.append("【个人简介】")
            lines.append(resume_data["summary"])
            lines.append("")

        for work in resume_data.get("work", []):
            lines.append("【工作经历】")
            lines.append(f"{work.get('company', '')} - {work.get('position', '')}")
            lines.append(self.format_date_range(work.get("start_date", ""), work.get("end_date", "")))
            if work.get("description"):
                lines.append(work["description"])
            for highlight in work.get("highlights", []):
                lines.append(f"• {highlight}")
            lines.append("")

        for edu in resume_data.get("education", []):
            lines.append("【教育背景】")
            lines.append(f"{edu.get('school', '')} - {edu.get('degree', '')} | {edu.get('major', '')}")
            lines.append(self.format_date_range(edu.get("start_date", ""), edu.get("end_date", "")))
            for highlight in edu.get("highlights", []):
                lines.append(f"• {highlight}")
            lines.append("")

        for project in resume_data.get("projects", []):
            lines.append("【项目经历】")
            lines.append(f"{project.get('name', '')} - {project.get('role', '')}")
            lines.append(self.format_date_range(project.get("start_date", ""), project.get("end_date", "")))
            if project.get("description"):
                lines.append(project["description"])
            for highlight in project.get("highlights", []):
                lines.append(f"• {highlight}")
            if project.get("technologies"):
                lines.append(f"技术栈：{', '.join(project['technologies'])}")
            lines.append("")

        if resume_data.get("skills"):
            lines.append("【专业技能】")
            for skill_category in resume_data["skills"]:
                lines.append(f"{skill_category.get('category', '')}：{', '.join(skill_category.get('items', []))}")
            lines.append("")

        content = "\n".join(lines)
        filename = f"{basics.get('name', 'resume')}_{resume_id}_{datetime.now().strftime('%Y%m%d')}.txt"
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return filepath

    def export_json(self, resume_data: Dict[str, Any], resume_id: str) -> str:
        filename = f"resume_{resume_id}_{datetime.now().strftime('%Y%m%d')}.json"
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(resume_data, f, ensure_ascii=False, indent=2)
        return filepath


export_service = ExportService()